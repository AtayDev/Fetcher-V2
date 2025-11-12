import os, hashlib, time
from pathlib import Path
from typing import List, Tuple, Dict

# Added send_from_directory
from flask import (
    Flask, render_template, request, jsonify, redirect, url_for, flash, send_from_directory
)
from werkzeug.utils import secure_filename
from dotenv import load_dotenv

import pdfplumber
import docx
import tiktoken

# LangChain imports
from langchain_community.vectorstores import Chroma
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain_core.documents import Document
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser

from datetime import datetime


def tprint(*args, **kwargs):
    print(f"[{datetime.now().strftime('%H:%M:%S')}] ", *args, **kwargs)


# =========================
# Config
# =========================
load_dotenv()

CHROMA_DIR = os.getenv("CHROMA_DIR", "./storage/chroma")
UPLOAD_DIR = os.getenv("UPLOAD_DIR", "./storage/uploads")
FLASK_SECRET = os.getenv("FLASK_SECRET", "dev-secret")
COLLECTION_NAME = os.getenv("COLLECTION_NAME", "global_knowledge")

LLM_PROVIDER = (os.getenv("LLM_PROVIDER") or "gemini").lower()

# Provider keys and models
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
OPENAI_EMBED_MODEL = os.getenv("OPENAI_EMBED_MODEL", "text-embedding-3-small")
OPENAI_CHAT_MODEL = os.getenv("OPENAI_CHAT_MODEL", "gpt-4o-mini")

GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY", "")
GEMINI_EMBED_MODEL = os.getenv("GEMINI_EMBED_MODEL", "text-embedding-004")
GEMINI_CHAT_MODEL = os.getenv("GEMINI_CHAT_MODEL", "gemini-1.5-flash")

DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY", "")
DEEPSEEK_BASE_URL = os.getenv("DEEPSEEK_BASE_URL", "https://api.deepseek.com")
DEEPSEEK_EMBED_MODEL = os.getenv("DEEPSEEK_EMBED_MODEL", "deepseek-embedding")
DEEPSEEK_CHAT_MODEL = os.getenv("DEEPSEEK_CHAT_MODEL", "deepseek-chat")

# Chunking / RAG settings
CHUNK_TOKENS = int(os.getenv("CHUNK_TOKENS", "600"))
CHUNK_OVERLAP = int(os.getenv("CHUNK_OVERLAP", "120"))
TOP_K = int(os.getenv("TOP_K", "4"))
FETCH_K = int(os.getenv("FETCH_K", "20"))
MMR_LAMBDA = float(os.getenv("MMR_LAMBDA", "0.4"))
CONTEXT_MAX_TOKENS = int(os.getenv("CONTEXT_MAX_TOKENS", "1200"))
ANSWER_MAX_TOKENS = int(os.getenv("ANSWER_MAX_TOKENS", "256"))

Path(CHROMA_DIR).mkdir(parents=True, exist_ok=True)
Path(UPLOAD_DIR).mkdir(parents=True, exist_ok=True)
ALLOWED_EXTS = {".pdf", ".docx", ".txt"}

app = Flask(__name__)
app.secret_key = FLASK_SECRET
_enc = tiktoken.get_encoding("cl100k_base")


# =========================
# Embedding + Chat factories
# =========================
def make_embeddings():
    if LLM_PROVIDER == "openai":
        return OpenAIEmbeddings(api_key=OPENAI_API_KEY, model=OPENAI_EMBED_MODEL)
    elif LLM_PROVIDER == "gemini":
        model = GEMINI_EMBED_MODEL if GEMINI_EMBED_MODEL.startswith("models/") else f"models/{GEMINI_EMBED_MODEL}"
        return GoogleGenerativeAIEmbeddings(google_api_key=GOOGLE_API_KEY, model=model)
    elif LLM_PROVIDER == "deepseek":
        return OpenAIEmbeddings(api_key=DEEPSEEK_API_KEY, model=DEEPSEEK_EMBED_MODEL, base_url=f"{DEEPSEEK_BASE_URL}/v1")
    raise RuntimeError(f"Unsupported provider {LLM_PROVIDER}")


def make_chat():
    if LLM_PROVIDER == "openai":
        return ChatOpenAI(api_key=OPENAI_API_KEY, model=OPENAI_CHAT_MODEL, temperature=0.1)
    elif LLM_PROVIDER == "gemini":
        return ChatGoogleGenerativeAI(google_api_key=GOOGLE_API_KEY, model=GEMINI_CHAT_MODEL, temperature=0.1)
    elif LLM_PROVIDER == "deepseek":
        return ChatOpenAI(api_key=DEEPSEEK_API_KEY, base_url=f"{DEEPSEEK_BASE_URL}/v1",
                          model=DEEPSEEK_CHAT_MODEL, temperature=0.1)
    raise RuntimeError(f"Unsupported provider {LLM_PROVIDER}")


# Lazy load
_EMB, _LLM = None, None
def EMB():
    global _EMB
    if _EMB is None: _EMB = make_embeddings()
    return _EMB

def LLM():
    global _LLM
    if _LLM is None: _LLM = make_chat()
    return _LLM


# =========================
# File utilities
# =========================
def allowed_file(fn: str) -> bool:
    return Path(fn).suffix.lower() in ALLOWED_EXTS


def file_sha256(p: Path) -> str:
    h = hashlib.sha256()
    with p.open("rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()


def read_pdf(path: Path) -> Tuple[str, List[str]]:
    pages = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            t = page.extract_text() or ""
            pages.append(t)
    return "\n".join(pages), pages


def read_docx(path: Path) -> str:
    d = docx.Document(str(path))
    return "\n".join([p.text for p in d.paragraphs if p.text.strip()])


def read_txt(path: Path) -> str:
    return path.read_text(errors="ignore")


def split_tokens(text: str, max_tokens=1200, overlap=250) -> List[str]:
    toks = _enc.encode(text)
    chunks, i, n = [], 0, len(toks)
    while i < n:
        j = min(n, i + max_tokens)
        seg = _enc.decode(toks[i:j]).strip()
        if seg:
            chunks.append(seg)
        i = max(0, j - overlap)
        if j >= n: break
    return chunks


# =========================
# Vectorstore (global)
# =========================
def vectorstore() -> Chroma:
    return Chroma(
        collection_name=COLLECTION_NAME,
        embedding_function=EMB(),
        persist_directory=CHROMA_DIR
    )


def add_documents(vs: Chroma, docs: List[Document], batch_size: int = 64):
    for start in range(0, len(docs), batch_size):
        batch = docs[start:start + batch_size]
        ids = [f"{d.metadata['sha']}::{int(time.time()*1e6)}::{i}" for i, d in enumerate(batch)]
        vs.add_documents(batch, ids=ids)
    vs.persist()


# =========================
# Chunking per file type
# =========================
def chunk_file(path: Path, filename: str, sha: str) -> List[Document]:
    suffix = path.suffix.lower()
    docs = []
    if suffix == ".pdf":
        _, pages = read_pdf(path)
        for pageno, text in enumerate(pages, 1): # Page numbers start at 1
            for i, chunk in enumerate(split_tokens(text, CHUNK_TOKENS, CHUNK_OVERLAP)):
                docs.append(Document(page_content=chunk, metadata={"filename": filename, "sha": sha,
                                                                  "page_start": pageno, "page_end": pageno,
                                                                  "chunk_index": i}))
    elif suffix == ".docx":
        text = read_docx(path)
        for i, chunk in enumerate(split_tokens(text, CHUNK_TOKENS, CHUNK_OVERLAP)):
            docs.append(Document(page_content=chunk, metadata={"filename": filename, "sha": sha,
                                                              "page_start": 1, "page_end": 1, # Use 1 for non-pdf
                                                              "chunk_index": i}))
    elif suffix == ".txt":
        text = read_txt(path)
        for i, chunk in enumerate(split_tokens(text, CHUNK_TOKENS, CHUNK_OVERLAP)):
            docs.append(Document(page_content=chunk, metadata={"filename": filename, "sha": sha,
                                                              "page_start": 1, "page_end": 1, # Use 1 for non-pdf
                                                              "chunk_index": i}))
    else:
        # Handle other text-based files if necessary, or raise error
        raise ValueError(f"Unsupported file type for chunking: {suffix}")
    
    return docs


# =========================
# Prompt / RAG
# =========================

SYSTEM_MESSAGE = """
You are an expert research assistant named "Nordlys Fetcher". Your goal is to answer the user's question based *only* on the provided context.

Follow these steps:
1.  **Reasoning:** First, think step-by-step about how to answer the question using the context. Analyze which parts of the context are most relevant, how they fit together, and what information is missing.
2.  **Synthesis:** Second, synthesize your findings into a comprehensive, well-structured answer.
3.  **Formatting:** Format your final answer using Markdown (e.g., `### Heading`, `**bold**`, `*italic*`, `- Bullet point`).
4.  **Tone:** Be professional, clear, and concise.

*Important:* Do not include your reasoning steps in the final output. Provide *only* the final, synthesized answer to the user.
"""

PROMPT = ChatPromptTemplate.from_messages([
    ("system", SYSTEM_MESSAGE),
    ("human", "Question:\n{question}\n\nContext:\n{context}\n\nAnswer:")
])


def trim_to_token_limit(texts: List[str], max_total_tokens: int):
    kept, used = [], 0
    for t in texts:
        nt = len(_enc.encode(t))
        if used + nt > max_total_tokens:
            break
        kept.append(t)
        used += nt
    return kept


def blocks_from_docs(docs: List[Document]) -> List[str]:
    blocks = []
    for doc in docs:
        md = doc.metadata
        fn = md.get("filename", "?")
        pg = md.get("page_start", -1)
        blk = f"{doc.page_content}\n[Source: {fn}, page {pg}]"
        blocks.append(blk)
    return blocks


# --- MODIFIED: Added doc_filter parameter ---
def rag_answer(question: str, doc_filter: List[str]) -> Dict:
    vs = vectorstore()
    
    # Base search arguments
    search_kwargs = {
        "k": TOP_K, 
        "fetch_k": FETCH_K, 
        "lambda_mult": MMR_LAMBDA
    }
    
    # --- NEW: Add document filter if provided ---
    if doc_filter:
        search_kwargs["filter"] = {"filename": {"$in": doc_filter}}
        tprint(f"Querying with filter: {doc_filter}")
    else:
        tprint("Querying all documents")
    
    retriever = vs.as_retriever(
        search_type="mmr",
        search_kwargs=search_kwargs
    )
    
    docs = retriever.invoke(question)
    if not docs:
        answer = "No relevant information found."
        if doc_filter:
            answer += f" in the selected documents: {', '.join(doc_filter)}"
        return {"answer": answer, "sources": []}

    blocks = blocks_from_docs(docs)
    context_text = "\n\n---\n\n".join(trim_to_token_limit(blocks, CONTEXT_MAX_TOKENS))
    chain = PROMPT | LLM() | StrOutputParser()
    answer = chain.invoke({"question": question, "context": context_text}).strip()
    
    sources = []
    seen = set()
    for doc in docs:
        md = doc.metadata
        fn = md.get("filename")
        pg = md.get("page_start", 1) 
        key = (fn, pg)
        is_pdf = fn.lower().endswith(".pdf")
        if fn and key not in seen:
            page_num = pg if is_pdf else 1
            sources.append({"filename": fn, "page": page_num, "is_pdf": is_pdf})
            seen.add(key)
            
    return {"answer": answer, "sources": sources[:5]}

"""
Time-limited token authentication system.
Add this to your app.py before the routes section.

Tokens format: base64(token_secret:timestamp:expiry_hours)
Example: dGVzdDoxNzMwNjQ2MDAwOjI0  (test token, valid for 24h)
"""

import base64
import hmac
import hashlib
from datetime import datetime, timedelta
from functools import wraps
from flask import request, abort, session, render_template, redirect, url_for, flash

# Secret key for signing tokens (MUST be in .env)
TOKEN_SECRET = os.getenv("TOKEN_SECRET", "change-me-in-production")

# Token validity period in hours (for testing: 30 seconds = 30/3600 hours)
TOKEN_EXPIRY_HOURS = float(os.getenv("TOKEN_EXPIRY_HOURS", str(72)))


def generate_token(client_id: str, expiry_hours: float = None) -> str:
    """
    Generate a time-limited token for a client.
    
    Args:
        client_id: Unique identifier for the client (e.g., "client1")
        expiry_hours: Hours until token expires (default: TOKEN_EXPIRY_HOURS)
    
    Returns:
        Base64-encoded token string
    """
    if expiry_hours is None:
        expiry_hours = TOKEN_EXPIRY_HOURS
    
    # Current timestamp
    issued_at = int(time.time())
    
    # Create token payload: client_id:timestamp:expiry_hours
    payload = f"{client_id}:{issued_at}:{expiry_hours}"
    
    # Sign the payload with HMAC
    signature = hmac.new(
        TOKEN_SECRET.encode(),
        payload.encode(),
        hashlib.sha256
    ).hexdigest()[:16]  # Use first 16 chars of signature
    
    # Combine payload and signature
    token_data = f"{payload}:{signature}"
    
    # Encode to base64
    token = base64.urlsafe_b64encode(token_data.encode()).decode()
    
    return token


def verify_token(token: str) -> dict:
    """
    Verify and decode a time-limited token.
    
    Returns:
        dict with 'valid', 'client_id', 'issued_at', 'expires_at', 'remaining_seconds'
        or dict with 'valid': False and 'error' message
    """
    try:
        # Decode from base64
        token_data = base64.urlsafe_b64decode(token.encode()).decode()
        
        # Split into parts
        parts = token_data.split(':')
        if len(parts) != 4:
            return {"valid": False, "error": "Invalid token format"}
        
        client_id, issued_at_str, expiry_hours_str, signature = parts
        
        # Reconstruct payload
        payload = f"{client_id}:{issued_at_str}:{expiry_hours_str}"
        
        # Verify signature
        expected_signature = hmac.new(
            TOKEN_SECRET.encode(),
            payload.encode(),
            hashlib.sha256
        ).hexdigest()[:16]
        
        if not hmac.compare_digest(signature, expected_signature):
            return {"valid": False, "error": "Invalid token signature"}
        
        # Parse timestamp and expiry
        issued_at = int(issued_at_str)
        expiry_hours = float(expiry_hours_str)
        
        # Calculate expiration
        issued_dt = datetime.fromtimestamp(issued_at)
        expires_at = issued_dt + timedelta(hours=expiry_hours)
        now = datetime.now()
        
        # Check if expired
        if now > expires_at:
            return {
                "valid": False,
                "error": "Token expired",
                "expired_at": expires_at.isoformat()
            }
        
        # Calculate remaining time
        remaining = (expires_at - now).total_seconds()
        
        return {
            "valid": True,
            "client_id": client_id,
            "issued_at": issued_dt.isoformat(),
            "expires_at": expires_at.isoformat(),
            "remaining_seconds": int(remaining),
            "remaining_human": format_time_remaining(remaining)
        }
    
    except Exception as e:
        return {"valid": False, "error": f"Token verification failed: {str(e)}"}


def format_time_remaining(seconds: float) -> str:
    """Format remaining seconds into human-readable string."""
    if seconds < 60:
        return f"{int(seconds)} seconds"
    elif seconds < 3600:
        minutes = int(seconds / 60)
        return f"{minutes} minute{'s' if minutes != 1 else ''}"
    else:
        hours = int(seconds / 3600)
        return f"{hours} hour{'s' if hours != 1 else ''}"


# Session timeout (for auto-logout)
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(hours=TOKEN_EXPIRY_HOURS)


@app.before_request
def check_token_auth():
    """Check authentication before every request."""
    
    # Public endpoints that don't require auth
    public_endpoints = ['static', 'login_token', 'logout']
    
    if request.endpoint in public_endpoints:
        return None
    
    # Check if already authenticated in session
    if session.get('authenticated'):
        # Verify session hasn't expired
        token_info = session.get('token_info')
        if token_info:
            expires_at = datetime.fromisoformat(token_info['expires_at'])
            if datetime.now() > expires_at:
                session.clear()
                flash('Your session has expired. Please log in again.', 'error')
                return redirect(url_for('login_token'))
        return None
    
    # Check for token in query params or header
    token = request.args.get('token') or request.headers.get('X-Access-Token')
    
    if token:
        # Verify token
        result = verify_token(token)
        
        if result['valid']:
            # Store in session
            session['authenticated'] = True
            session['token_info'] = result
            session.permanent = True
            
            tprint(f"✅ Client authenticated: {result['client_id']}, "
                   f"expires in {result['remaining_human']}")
            
            return None
        else:
            # Invalid or expired token
            tprint(f"❌ Authentication failed: {result.get('error')}")
            flash(f"Authentication failed: {result.get('error')}", 'error')
    
    # No valid authentication - redirect to login
    if request.endpoint == 'index' and request.method == 'GET':
        return render_template('login_token.html')

    if request.accept_mimetypes.accept_html:
        return redirect(url_for('login_token'))

    abort(401)


@app.route('/login')
def login_token():
    """Login page for token entry."""
    return render_template('login_token.html')


@app.route('/logout')
def logout():
    """Logout and clear session."""
    client_id = session.get('token_info', {}).get('client_id', 'Unknown')
    session.clear()
    tprint(f"🚪 Client logged out: {client_id}")
    flash('Logged out successfully', 'ok')
    return redirect(url_for('login_token'))


@app.route('/api/token/info')
def token_info():
    """Get current token information (for debugging)."""
    if not session.get('authenticated'):
        return jsonify({"error": "Not authenticated"}), 401
    
    info = session.get('token_info', {})
    return jsonify(info)

# =========================
# Routes
# =========================
@app.get("/")
def index():
    files = [p.name for p in Path(UPLOAD_DIR).glob("*") if p.is_file()]
    return render_template("index.html", files=files)


@app.get("/files/<path:filename>")
def get_file(filename):
    """Serves uploaded files from the UPLOAD_DIR."""
    return send_from_directory(UPLOAD_DIR, filename, as_attachment=False)


# --- NEW: API to list indexed files ---
@app.get("/api/files")
def api_files():
    """Returns a list of all indexed filenames."""
    try:
        files = [p.name for p in Path(UPLOAD_DIR).glob("*") if p.is_file()]
        return jsonify(files=files)
    except Exception as e:
        tprint(f"Error in /api/files: {e}")
        return jsonify({"error": str(e)}), 500


@app.post("/upload")
def upload():
    f = request.files.get("file")
    if not f or f.filename == "":
        flash("No file selected", "error")
        return redirect(url_for("index"))
    if not allowed_file(f.filename):
        flash("Invalid file type", "error")
        return redirect(url_for("index"))

    filename = secure_filename(f.filename)
    dest = Path(UPLOAD_DIR) / filename
    
    if dest.exists():
        flash(f"Error: File '{filename}' already exists.", "error")
        return redirect(url_for("index"))

    f.save(dest)
    sha = file_sha256(dest)
    tprint(f"Processing file {filename} (sha={sha[:8]}…)")

    try:
        docs = chunk_file(dest, filename, sha)
        vs = vectorstore()
        add_documents(vs, docs)
        flash(f"Uploaded & indexed: {filename} ({len(docs)} chunks)", "ok")
    except Exception as e:
        dest.unlink() # Clean up failed upload
        tprint(f"Error indexing file {filename}: {e}")
        flash(f"Error indexing file: {e}", "error")

    return redirect(url_for("index"))


@app.post("/api/chat")
def api_chat():
    data = request.get_json(silent=True) or {}
    q = data.get("q", "").strip()
    # --- NEW: Get document filter from request ---
    doc_filter = data.get("doc_filter", []) 
    
    if not q:
        return jsonify({"answer": "Please provide a question."})
    try:
        # --- NEW: Pass filter to rag_answer ---
        result = rag_answer(q, doc_filter)
        return jsonify(result)
    except Exception as e:
        tprint(f"Error in /api/chat: {e}")
        return jsonify({"answer": f"Error: {e}", "sources": []}), 500


@app.get("/api/stats")
def api_stats():
    """Returns statistics about the indexed documents."""
    try:
        files = [p for p in Path(UPLOAD_DIR).glob("*") if p.is_file()]
        doc_count = len(files)
        total_size = sum(p.stat().st_size for p in files)
        
        last_indexed_ts = 0
        if files:
            last_indexed_ts = max(p.stat().st_mtime for p in files)
        
        last_indexed_str = "N/A"
        if last_indexed_ts > 0:
            last_indexed_str = datetime.fromtimestamp(last_indexed_ts).strftime("%Y-%m-%d")

        # Format size
        if total_size > 1024 * 1024:
            size_str = f"{total_size / (1024 * 1024):.2f} MB"
        elif total_size > 1024:
            size_str = f"{total_size / 1024:.2f} KB"
        else:
            size_str = f"{total_size} B"

        return jsonify({
            "doc_count": doc_count,
            "total_size": size_str,
            "last_indexed": last_indexed_str
        })
    except Exception as e:
        tprint(f"Error in /api/stats: {e}")
        return jsonify({"error": str(e)}), 500


@app.post("/api/delete")
def api_delete():
    """Deletes a file and its vectors from the index."""
    data = request.get_json(silent=True) or {}
    filename = data.get("filename")
    
    if not filename:
        return jsonify({"error": "No filename provided"}), 400
    
    sec_filename = secure_filename(filename)
    if sec_filename != filename:
        return jsonify({"error": "Invalid filename"}), 400
        
    try:
        # 1. Delete file from disk
        file_path = Path(UPLOAD_DIR) / sec_filename
        if not file_path.exists() or not file_path.is_file():
            return jsonify({"error": "File not found on disk"}), 404
        
        file_path.unlink()
        tprint(f"Deleted file from disk: {sec_filename}")

        # 2. Delete vectors from ChromaDB
        vs = vectorstore()
        vs.delete(where={"filename": sec_filename})
        vs.persist()
        tprint(f"Deleted vectors from Chroma for: {sec_filename}")

        return jsonify({"success": True, "filename": sec_filename})

    except Exception as e:
        tprint(f"Error in /api/delete for {filename}: {e}")
        return jsonify({"error": str(e)}), 500


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=8000, debug=True)
